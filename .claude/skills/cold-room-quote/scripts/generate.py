#!/usr/bin/env python3
"""Generate a TNDK cold room quotation from the saved master template.

The master (template/master.docx) is never rebuilt from scratch — it is copied
and edited in place, so the cover page, brand colours, header strip, footer and
all thirteen section illustrations survive untouched.

Usage:
    python3 generate.py --spec quote.json --output /path/to/quotation.docx
    python3 generate.py --spec quote.json --output out.docx --print-summary

See references/spec-schema.md for the full JSON shape and
references/template-map.md for what lives where inside the document.
"""

import argparse
import copy
import json
import re
import shutil
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from number_to_words import amount_in_words  # noqa: E402

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

SKILL_DIR = Path(__file__).resolve().parent.parent
MASTER = SKILL_DIR / "template" / "master.docx"

# Prose strings as they appear in the frozen master. Editing the master means
# updating these — nothing else in the script hardcodes document text.
MASTER_INTRO = (
    "Dear Sir, with reference to the above subject, we thank you for the "
    "opportunity to quote for your cold storage room project. Please find below "
    "our detailed technical and commercial offer for the supply, installation, "
    "testing and commissioning of inclusive of insulated flooring, refrigeration "
    "system and accessories."
)
AMOUNT_PREFIX = "Amount in Words:"

# Doha design ambient, used when a quote does not state one
DEFAULT_AMBIENT = "46°C"

# Height for the service banner under section 12. The free space below the
# delivery table varies with how section 10 breaks — measured at 5.3 in on a
# one-room quote and 4.4 in on a two-room one — and the banner needs its cell
# padding on top of its own height. 4.0 in clears the tighter case; lower
# banner_height_in if a quotation still strands it on a page of its own.
BANNER_HEIGHT_IN = 4.0


# --------------------------------------------------------------------------
# low-level docx helpers — all of these preserve existing run formatting
# --------------------------------------------------------------------------

def para_text(p):
    return "".join(r.text for r in p.runs)


def set_para_text(p, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def replace_in_para(p, old, new):
    """Replace `old` inside a paragraph even when it spans several runs.

    The replacement inherits the formatting of the run the match starts in, so
    mixed-format lines (bold blue label + black value) keep their look.
    """
    runs = p.runs
    full = "".join(r.text for r in runs)
    if old not in full:
        return False
    start = full.index(old)
    end = start + len(old)
    pos = 0
    first = True
    for r in runs:
        r_start, r_end = pos, pos + len(r.text)
        pos = r_end
        if r_end <= start or r_start >= end:
            continue
        head = r.text[: max(0, start - r_start)]
        tail = r.text[max(0, end - r_start):]
        if first:
            r.text = head + new + tail
            first = False
        else:
            r.text = head + tail
    return True


def set_cell(cell, value):
    """Write `value` into a table cell, one paragraph per line."""
    lines = value.split("\n") if isinstance(value, str) else list(value)
    paras = cell.paragraphs
    for i, line in enumerate(lines):
        if i < len(paras):
            set_para_text(paras[i], line)
        else:
            new_p = copy.deepcopy(paras[0]._p)
            paras[0]._p.getparent().append(new_p)
            cell_para = cell.paragraphs[-1]
            set_para_text(cell_para, line)
    # blank out any surplus paragraphs the master had
    for p in cell.paragraphs[len(lines):]:
        set_para_text(p, "")


def row_cells(row):
    """Cells of a row, de-duplicated across horizontal merges."""
    out, seen = [], set()
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            out.append(c)
    return out


def find_table(doc, *labels):
    """Find the table whose first column contains all of `labels`."""
    for t in doc.tables:
        firsts = {row_cells(r)[0].text.strip() for r in t.rows}
        if all(lb in firsts for lb in labels):
            return t
    raise LookupError(f"no table with first-column labels {labels}")


def set_row(table, label, value):
    """Set the value cell of the row whose first cell is `label`."""
    for r in table.rows:
        cells = row_cells(r)
        if cells[0].text.strip() == label:
            set_cell(cells[-1], value)
            return True
    raise LookupError(f"row {label!r} not found")


def clone_row(table, template_row_idx=-1):
    src = table.rows[template_row_idx]._tr
    new = copy.deepcopy(src)
    src.getparent().append(new)
    return table.rows[-1]


# --------------------------------------------------------------------------
# quantity calculations
# --------------------------------------------------------------------------

def fmt(n, dp=2):
    """3.0 -> '3', 14.4 -> '14.40' — matches how the master writes numbers."""
    n = float(n)
    if abs(n - round(n)) < 1e-9 and dp == 0:
        return str(int(round(n)))
    return f"{n:.{dp}f}"


def dim(n):
    """Dimension exactly as it should read on the page.

    Written as a string in the spec ("2.40") it is reproduced verbatim, because
    JSON numbers cannot carry a trailing zero and the master writes 2.40, not
    2.4. Plain numbers fall back to at most two decimals, at least one.
    """
    if isinstance(n, str):
        return n.strip()
    s = f"{float(n):.2f}".rstrip("0")
    return s + "0" if s.endswith(".") else s


def room_quantities(room):
    """Panel coverage, floor area and internal volume for one room."""
    L, W, H = float(room["length"]), float(room["width"]), float(room["height"])
    dL, dW, dH = dim(room["length"]), dim(room["width"]), dim(room["height"])
    qty = int(room.get("qty", 1))
    walls = (2 * (L * H) + 2 * (W * H)) * qty
    ceiling = L * W * qty
    floor = L * W * qty if room.get("floor_included", True) else 0.0
    volume = L * W * H * qty
    return {
        "L": L, "W": W, "H": H, "qty": qty,
        "dL": dL, "dW": dW, "dH": dH,
        "walls": walls, "ceiling": ceiling, "floor": floor,
        "total_panel": walls + ceiling + floor,
        "volume": volume,
        "wall_expr": f"2 × ({dL} × {dH}) + 2 × ({dW} × {dH})",
        "area_expr": f"{dL} × {dW}",
        "vol_expr": f"{dL} × {dW} × {dH}",
        "dims": f"{dL} × {dW} × {dH} m",
    }


def totals(rooms):
    qs = [room_quantities(r) for r in rooms]
    return qs, {
        "walls": sum(q["walls"] for q in qs),
        "ceiling": sum(q["ceiling"] for q in qs),
        "floor": sum(q["floor"] for q in qs),
        "total_panel": sum(q["total_panel"] for q in qs),
        "volume": sum(q["volume"] for q in qs),
    }


def multiplied(q, per_room_expr, value):
    """'2 × (6.0 × 2.90) + ... = 48.72 sqm', with a ×N when qty > 1."""
    suffix = f" × {q['qty']} = {fmt(value)} sqm" if q["qty"] > 1 else f" = {fmt(value)} sqm"
    return per_room_expr + suffix


# --------------------------------------------------------------------------
# document population
# --------------------------------------------------------------------------

def fill_reference_block(doc, spec):
    t = find_table(doc, "REFERENCE")
    cells = row_cells(t.rows[1])
    set_cell(cells[0], spec["reference"])
    set_cell(cells[1], spec["date"])
    set_cell(cells[2], spec.get("validity", "15 Days from quote date"))


def fill_client_block(doc, spec):
    t = find_table(doc, "TO")
    client = spec["client"]
    lines = [client["name"], ""]
    lines += [ln for ln in str(client.get("address", "")).split("\n") if ln]
    if client.get("attn"):
        lines.append(f"Attn: {client['attn']}")
    if client.get("phone"):
        lines.append(f"Tel: {client['phone']}")
    set_cell(row_cells(t.rows[1])[0], lines)


def fill_project_table(doc, spec, qs):
    t = find_table(doc, "ROOM TYPE")
    # master ships with exactly one data row; add more when quoting more rooms
    while len(t.rows) - 1 < len(spec["rooms"]):
        clone_row(t)
    for i, (room, q) in enumerate(zip(spec["rooms"], qs), start=1):
        cells = row_cells(t.rows[i])
        set_cell(cells[0], room["type"].upper())
        set_cell(cells[1], str(q["qty"]))
        set_cell(cells[2], room["temperature"])
        set_cell(cells[3], q["dims"])
        set_cell(cells[4], room.get("panel_thickness", "100 mm"))


def fill_panel_section(doc, spec, qs, tot):
    t = find_table(doc, "Panel Type", "Wall Coverage")
    set_row(t, "Panel Thickness", thickness_label(spec))
    if len(qs) == 1:
        q = qs[0]
        set_row(t, "Wall Coverage", multiplied(q, q["wall_expr"], q["walls"]))
        set_row(t, "Ceiling Coverage", multiplied(q, q["area_expr"], q["ceiling"]))
        set_row(t, "Floor Coverage",
                multiplied(q, q["area_expr"], q["floor"]) if q["floor"]
                else "Not included in scope")
    else:
        set_row(t, "Wall Coverage", f"{fmt(tot['walls'])} sqm (all rooms)")
        set_row(t, "Ceiling Coverage", f"{fmt(tot['ceiling'])} sqm (all rooms)")
        set_row(t, "Floor Coverage",
                f"{fmt(tot['floor'])} sqm (all rooms)" if tot["floor"]
                else "Not included in scope")
    set_row(t, "Total Panel Quantity", f"{fmt(tot['total_panel'])} sqm")


def door_kind(spec):
    """'hinged' or 'sliding' — the master is written for hinged doors."""
    return str(spec.get("door", {}).get("type", "hinged")).strip().lower()


def fill_door_section(doc, spec):
    t = find_table(doc, "Door Type")
    door = spec.get("door", {})
    w = door.get("width_mm", 900)
    h = door.get("height_mm", 1900)
    kind = door_kind(spec)
    set_row(t, "Clear Opening Size", f"{w} mm (W) × {h} mm (H)")
    set_row(t, "Door Thickness", f"{thickness_label(spec)} — matching wall panel")
    set_row(t, "Quantity", str(door.get("qty", 1)))
    if kind == "sliding":
        # the master describes a hinged leaf; a sliding door hangs on a track
        # and has no hinges, so those rows would otherwise be plainly wrong
        set_row(t, "Door Type", "Sliding Cold Room Door (single leaf)")
        for r in t.rows:
            cells = row_cells(r)
            if cells[0].text.strip() == "Hinges":
                set_cell(cells[0], "Track & Rollers")
                set_cell(cells[-1], "Heavy-duty overhead track with nylon rollers, "
                                    "guide and end stops")
                break
        set_row(t, "Handle / Lock",
                "Recessed sliding handle with internal safety release (anti-trap)")
    else:
        set_row(t, "Door Type", "Hinged Cold Room Door (single leaf)")


NO_FLOOR = "Not included in this offer"


def fill_flooring_section(doc, spec, tot):
    t = find_table(doc, "Insulated Floor Panel")
    if not tot["floor"]:
        set_row(t, "Insulated Floor Panel",
                f"{NO_FLOOR} — room erected on the client's existing finished floor")
        set_row(t, "Plywood Layer", NO_FLOOR)
        set_row(t, "Top Finish", NO_FLOOR)
        set_row(t, "Edge Sealing",
                "Silicone seal between the wall panels and the existing floor")
        return
    area = f"{fmt(tot['floor'])} sqm"
    floored = [r for r in spec["rooms"] if r.get("floor_included", True)]
    thickness = thickness_label({"rooms": floored or spec["rooms"]})
    set_row(t, "Insulated Floor Panel",
            f"{thickness} PUF panel, same specification as wall panels — {area}")
    set_row(t, "Plywood Layer",
            f"18 mm Marine-grade plywood laid over the floor panel — {area}")
    set_row(t, "Top Finish",
            f"3 mm Mild Steel Chequered (Checkered) Plate, slip-resistant — {area}")


def ref_entries(spec):
    """Refrigeration as a list — one entry per system, however it was written.

    A single-system quote may pass one object; a job with a different machine
    per room passes a list, each entry optionally carrying a `room` label that
    prefixes its line in the tables.
    """
    ref = spec.get("refrigeration")
    if not ref:
        return []
    return list(ref) if isinstance(ref, list) else [ref]


def _tag(entry):
    return f"{entry['room']}: " if entry.get("room") else ""


def _origin(entry, key):
    return f" ({entry[key]})" if entry.get(key) else ""


def cu_text(entry):
    kind = entry.get("cu_type", "semi-hermetic").strip()
    # "BITZER LH64/2DES-3Y, semi-hermetic condensing unit" reads with a comma;
    # a bare brand does not — "Copeland condensing unit"
    head = (f"{entry['condensing_unit']}, {kind} condensing unit" if kind
            else f"{entry['condensing_unit']} condensing unit")
    return head + (f", {entry['cu_origin']}" if entry.get("cu_origin") else "")


def compressor_text(entry):
    brand = entry.get("compressor_brand") or entry["condensing_unit"].split()[0]
    # the model is quoted without repeating the brand that already precedes it
    model = re.sub(rf"^{re.escape(brand)}\s*", "", entry["condensing_unit"]).strip()
    kind = entry.get("cu_type", "semi-hermetic").strip()
    text = f"{brand}{_origin(entry, 'cu_origin')}"
    if kind:
        text += f" — {kind} compressor"
    if model:
        text += f"{',' if kind else ' —'} model {model}"
    return text


def evaporator_text(entry):
    parts = [f"{entry['evaporator']} evaporating unit"]
    if entry.get("evap_note"):
        parts.append(entry["evap_note"])
    if entry.get("evap_origin"):
        parts.append(entry["evap_origin"])
    return ", ".join(parts)


def capacity_of(entry):
    """'5.60 kW', or whatever capacity_label says when a unit is sold by HP."""
    if entry.get("capacity_label"):
        return entry["capacity_label"]
    return f"{fmt(entry['capacity_kw'])} kW"


def capacity_text(entry):
    basis = entry.get("selection_basis") or (
        entry.get("compressor_brand")
        or entry.get("condensing_unit", "").split()[0] if entry.get("condensing_unit") else "")
    at = f"{entry['sst']} SST / {entry['ambient']}" if entry.get("sst") else entry.get("ambient", "")
    note = f" (per {basis} selection at {at} ambient)" if basis and at else ""
    return f"Cooling capacity {capacity_of(entry)}{note}"


def selected_text(entry):
    bits = [f"{capacity_of(entry)} cooling capacity"]
    if entry.get("power_input_kw"):
        bits.append(f"{fmt(entry['power_input_kw'])} kW power input")
    if entry.get("current_a"):
        supply = entry.get("power_supply", "400V-3-50Hz")
        bits.append(f"{fmt(entry['current_a'])} A @ {supply}")
    head = entry.get("condensing_unit", "")
    return (f"{head} — " if head else "") + ", ".join(bits)


def joined(entries, fn):
    """One labelled line per system — collapsed to one line when they agree.

    Two rooms sharing a brand should not print that brand twice; two rooms with
    different capacities should print both, each tagged with its room.
    """
    texts = [fn(e) for e in entries]
    if len(set(texts)) == 1:
        return texts[0]
    return "\n".join(_tag(e) + t for e, t in zip(entries, texts))


def fill_refrigeration_section(doc, spec):
    entries = ref_entries(spec)
    if not entries:
        return
    t = find_table(doc, "System Type")
    sets_ = spec.get("sets") or entries[0].get("sets") or len(entries)
    set_row(t, "System Type", f"Split-type refrigeration system "
            f"({sets_} set{'s' if sets_ > 1 else ''})")
    if all(e.get("condensing_unit") for e in entries):
        set_row(t, "Condensing Unit", joined(entries, cu_text))
        set_row(t, "Compressor Brand", joined(entries, compressor_text))
    refrigerants = []
    for e in entries:
        if e.get("refrigerant") and e["refrigerant"] not in refrigerants:
            refrigerants.append(e["refrigerant"])
    if refrigerants:
        set_row(t, "Refrigerant", " / ".join(refrigerants))
    if all(e.get("evaporator") for e in entries):
        set_row(t, "Evaporator (Air Cooler)", joined(entries, evaporator_text))
    set_row(t, "Operating Temperature", operating_temp(spec))
    defrosts = [e["defrost"] for e in entries if e.get("defrost")]
    if defrosts:
        set_row(t, "Defrost", defrosts[0])


def operating_temp(spec):
    """'0°C to +8°C (Chiller)', or one clause per room when they differ."""
    parts, seen = [], set()
    for room in spec["rooms"]:
        clause = f"{room['temperature']} ({room['type'].title()})"
        if clause not in seen:
            seen.add(clause)
            parts.append(clause)
    return " / ".join(parts)


def thickness_label(spec):
    """'100 mm', or '100 / 150 mm' when rooms use different panels."""
    seen = []
    for room in spec["rooms"]:
        t = str(room.get("panel_thickness", "100 mm")).strip()
        if t not in seen:
            seen.append(t)
    if len(seen) == 1:
        return seen[0]
    numbers = [t.replace("mm", "").strip() for t in seen]
    return " / ".join(numbers) + " mm"


def fill_capacity_section(doc, spec, tot, qs):
    t = find_table(doc, "Internal Volume")
    entries = ref_entries(spec)
    if len(qs) == 1:
        vol = f"{qs[0]['vol_expr']} = {fmt(tot['volume'])} m³"
    else:
        vol = "\n".join(
            f"{r['type'].title()}: {q['vol_expr']} = {fmt(q['volume'])} m³"
            for r, q in zip(spec["rooms"], qs)) + f"\nTotal: {fmt(tot['volume'])} m³"
    set_row(t, "Internal Volume", vol)
    set_row(t, "Design Room Temp.", operating_temp(spec))

    rated = [e for e in entries if e.get("capacity_kw") or e.get("capacity_label")]
    if rated:
        set_row(t, "Estimated Heat Load", joined(rated, capacity_text))
        set_row(t, "Selected Capacity", joined(rated, selected_text))
    if entries:
        ambients = []
        for e in entries:
            if e.get("ambient") and e["ambient"] not in ambients:
                ambients.append(e["ambient"])
        # Always rewrite this row once machines are quoted: leaving the master's
        # value would carry the reference quote's ambient and its brand into a
        # quotation selected on entirely different equipment.
        basis = entries[0].get("selection_basis") if len(entries) == 1 else None
        note = f"per {basis} selection, " if basis else ""
        set_row(t, "Design Ambient",
                f"+{(ambients[0] if ambients else DEFAULT_AMBIENT).lstrip('+')} "
                f"({note}Doha ambient design)")


def fill_scope(doc, spec, tot):
    """Rewrite the scope bullets whose wording depends on the quoted sizes."""
    thickness = thickness_label(spec)
    door = spec.get("door", {})
    entries = ref_entries(spec)
    sets_ = spec.get("sets") or (entries[0].get("sets") if entries else None) \
        or max(len(entries), 1)
    n_words = {1: "one (1)", 2: "two (2)", 3: "three (3)", 4: "four (4)"}
    door_qty = door.get("qty", 1)
    kind = door_kind(spec)
    surfaces = "walls, ceiling and floor" if tot["floor"] else "walls and ceiling"
    floor_bullet = (
        f"Supply and installation of insulated floor system: {thickness} PUF floor "
        f"panel + 18 mm marine plywood + 3 mm MS chequered plate."
        if tot["floor"] else
        "Insulated flooring is not included in this offer — the room will be "
        "erected on the client's existing finished floor.")
    rules = [
        (r"^Supply of .* sandwich panels for walls.*$",
         f"Supply of {thickness} PUF sandwich panels for {surfaces} "
         f"(total {fmt(tot['total_panel'])} sqm)."),
        (r"^Supply and installation of .* hinged cold room door.*$",
         f"Supply and installation of {n_words.get(door_qty, str(door_qty))} No. "
         f"{kind} cold room door{'s' if door_qty > 1 else ''} "
         f"({door.get('width_mm', 900)} × "
         f"{door.get('height_mm', 1900)} mm)."),
        (r"^Supply and installation of insulated floor system.*$", floor_bullet),
        (r"^Supply and installation of condensing unit.*$",
         f"Supply and installation of condensing unit and matching evaporator "
         f"({sets_} set{'s' if sets_ > 1 else ''})."),
    ]
    for p in doc.paragraphs:
        text = para_text(p).strip()
        for pattern, replacement in rules:
            if re.match(pattern, text):
                set_para_text(p, replacement)
                break


def fill_boq(doc, spec, qs, tot):
    t = find_table(doc, "SL#")
    items = spec.get("boq")
    if not items:
        items = default_boq(spec, qs, tot)
    while len(t.rows) - 1 < len(items):
        clone_row(t)
    for i, item in enumerate(items, start=1):
        cells = row_cells(t.rows[i])
        set_cell(cells[0], str(i))
        set_cell(cells[1], item["description"] if isinstance(item, dict) else str(item))
        amount = item.get("amount") if isinstance(item, dict) else None
        set_cell(cells[2], f"{float(amount):,.2f}" if amount else "")
    # blank any rows the master had beyond what this quote needs
    for r in t.rows[len(items) + 1:]:
        for c in row_cells(r):
            set_cell(c, "")


def default_boq(spec, qs, tot):
    thickness = thickness_label(spec)
    door = spec.get("door", {})
    entries = ref_entries(spec)
    q = qs[0]
    if len(qs) == 1:
        size = (f"internal size {q['dL']} m (L) × {q['dW']} m (W) × "
                f"{q['dH']} m (H) ({fmt(q['volume'])} m³)")
    else:
        size = " and ".join(
            f"{r['type'].lower()} {x['dims']} ({fmt(x['volume'])} m³)"
            for r, x in zip(spec["rooms"], qs))

    def one_system(e):
        text = f"{e['condensing_unit']}"
        kind = e.get("cu_type", "semi-hermetic").strip()
        text += f" {kind} condensing unit" if kind else " condensing unit"
        text += _origin(e, "cu_origin")
        if e.get("evaporator"):
            text += f" + {e['evaporator']} evaporating unit{_origin(e, 'evap_origin')}"
        if e.get("refrigerant"):
            text += f", {e['refrigerant']}"
        if e.get("capacity_kw") or e.get("capacity_label"):
            text += f", cooling capacity {capacity_of(e)}"
        return _tag(e) + text

    fridge = ""
    if entries and all(e.get("condensing_unit") for e in entries):
        fridge = "; ".join(one_system(e) for e in entries)
        fridge += (", including refrigeration copper piping, armaflex insulation "
                   "and insulated drain line")
    n_words = {1: "one (1)", 2: "two (2)", 3: "three (3)", 4: "four (4)"}
    door_qty = door.get("qty", 1)
    door_plural = "s" if door_qty > 1 else ""
    kind = door_kind(spec)
    gear = ("heavy-duty track, rollers and safety handle with inside release"
            if kind == "sliding" else
            "heavy-duty hinges, safety latch handle and inside release")
    panel_area = (f" (total panel area incl. floor: {fmt(tot['total_panel'])} sqm)"
                  if tot["floor"] else
                  f" (total panel area: {fmt(tot['total_panel'])} sqm)")
    items = [
        {"description":
            f"Supply, fabrication & installation of cold room — {size}, complete with "
            f"{thickness} PUF sandwich panels for walls & ceiling{panel_area}, all "
            f"aluminum coving, corner angles, silicone sealing, fasteners and "
            f"accessories."},
        {"description":
            f"Supply & installation of {n_words.get(door_qty, str(door_qty))} No. {kind} "
            f"cold room door{door_plural}, clear opening size {door.get('width_mm', 900)} × "
            f"{door.get('height_mm', 1900)} mm, {thickness} thick, with {gear}."},
    ]
    if tot["floor"]:
        items.append({"description":
            f"Supply & installation of insulated floor system: {thickness} PUF floor panel "
            f"+ 18 mm marine plywood + 3 mm MS chequered plate ({fmt(tot['floor'])} sqm)."})
    items += [
        {"description": f"Supply & installation of refrigeration system: {fridge}."},
        {"description":
            "Supply & installation of digital control panel with all safeties, internal LED "
            "light, door frame heater and complete electrical works from control panel onward."},
        {"description":
            "Testing, commissioning, transportation, mobilization, lifting at site and "
            "handover with operation training."},
    ]
    return items


def fill_total(doc, spec):
    t = find_table(doc, "GRAND TOTAL (Lump Sum)")
    total = float(spec["total"])
    set_cell(row_cells(t.rows[0])[-1], f"QAR {total:,.2f}")
    words = spec.get("total_words") or amount_in_words(total)
    for p in doc.paragraphs:
        if para_text(p).strip().startswith(AMOUNT_PREFIX):
            current = para_text(p)[len(AMOUNT_PREFIX):].strip()
            replace_in_para(p, current, " " + words if not current else words)
            break


def fill_delivery(doc, spec):
    delivery = spec.get("delivery")
    if not delivery:
        return
    t = find_table(doc, "Material Delivery")
    for label, value in delivery.items():
        try:
            set_row(t, label, value)
        except LookupError:
            print(f"  ! delivery row {label!r} not in template — skipped",
                  file=sys.stderr)


def fit_banner(doc, spec):
    """Shrink the service banner so it sits under section 12 instead of
    stranding itself on a page of its own.

    The master ships the banner at 7.27 x 5.45 in, which is a shade taller than
    the space left below the delivery table — so it flows to the next page and
    the quotation gains a near-empty sheet. Scaling it to fit closes that gap.
    Set `banner_height_in` to 0 to leave the master's size alone.
    """
    target_h = spec.get("banner_height_in", BANNER_HEIGHT_IN)
    if not target_h:
        return
    body = doc.element.body
    # The master parks the banner in the delivery table's own trailing row, but
    # a revised master could just as well put it in a paragraph underneath.
    candidates, seen_delivery = [], False
    for child in body.iterchildren():
        if child.tag == qn("w:tbl"):
            # joined, because Word may split a label across several runs
            texts = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if "Material Delivery" in texts:
                seen_delivery = True
                candidates.append(child)
            continue
        if seen_delivery and child.tag == qn("w:p"):
            candidates.append(child)

    for child in candidates:
        extent = child.find(".//" + qn("wp:extent"))
        if extent is None:
            continue
        cx, cy = int(extent.get("cx")), int(extent.get("cy"))
        target_cy = int(target_h * 914400)
        if cy <= target_cy:
            return
        target_cx = int(cx * target_cy / cy)
        extent.set("cx", str(target_cx))
        extent.set("cy", str(target_cy))
        for ext in child.iter(qn("a:ext")):
            # a:ext also names the extension-list element, which carries a uri
            # and rejects cx/cy — only the a:xfrm one describes a size
            if ext.get("cx") is None:
                continue
            ext.set("cx", str(target_cx))
            ext.set("cy", str(target_cy))
        return


def fill_subject(doc, spec, tot=None):
    subject = spec.get("subject")
    intro = spec.get("intro")
    if not intro and tot is not None and not tot["floor"]:
        # the master's wording promises insulated flooring, which is not supplied here
        # also repairs the master's "commissioning of inclusive of" slip
        intro = MASTER_INTRO.replace("of inclusive of insulated flooring, "
                                     "refrigeration system and accessories",
                                     "of the cold storage room, inclusive of the "
                                     "refrigeration system and accessories")
    for p in doc.paragraphs:
        text = para_text(p).strip()
        if subject and text.startswith("Subject:"):
            replace_in_para(p, text[len("Subject:"):].strip(), subject)
        elif intro and text == MASTER_INTRO:
            set_para_text(p, intro)


# --------------------------------------------------------------------------

def validate(spec):
    missing = [k for k in ("reference", "date", "client", "rooms", "total")
               if k not in spec]
    if missing:
        sys.exit(f"spec is missing required field(s): {', '.join(missing)}")
    if not spec["rooms"]:
        sys.exit("spec needs at least one room")
    for i, room in enumerate(spec["rooms"], 1):
        for k in ("type", "temperature", "length", "width", "height"):
            if k not in room:
                sys.exit(f"room {i} is missing '{k}'")
    if "name" not in spec["client"]:
        sys.exit("client needs a 'name'")
    if not re.match(r"^QUT/DCTS/[A-Z]{0,2}\d+/\d{4}$", spec["reference"]):
        print(f"  ! reference {spec['reference']!r} does not match "
              f"QUT/DCTS/[SQ]NNN/YYYY — continuing anyway", file=sys.stderr)


def generate(spec, output):
    validate(spec)
    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(MASTER, output)
    doc = docx.Document(output)

    qs, tot = totals(spec["rooms"])

    fill_reference_block(doc, spec)
    fill_client_block(doc, spec)
    fill_subject(doc, spec, tot)
    fill_project_table(doc, spec, qs)
    fill_panel_section(doc, spec, qs, tot)
    fill_door_section(doc, spec)
    fill_flooring_section(doc, spec, tot)
    fill_refrigeration_section(doc, spec)
    fill_capacity_section(doc, spec, tot, qs)
    fill_scope(doc, spec, tot)
    fill_boq(doc, spec, qs, tot)
    fill_total(doc, spec)
    fill_delivery(doc, spec)
    fit_banner(doc, spec)

    doc.save(output)
    return output, tot


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--spec", required=True, help="JSON spec file (see references/spec-schema.md)")
    ap.add_argument("--output", required=True, help="path for the generated .docx")
    ap.add_argument("--print-summary", action="store_true",
                    help="print computed quantities for cross-checking")
    args = ap.parse_args()

    spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
    output, tot = generate(spec, args.output)

    print(f"Written: {output}")
    if args.print_summary:
        print(f"  Reference   : {spec['reference']}")
        print(f"  Client      : {spec['client']['name']}")
        print(f"  Wall panel  : {fmt(tot['walls'])} sqm")
        print(f"  Ceiling     : {fmt(tot['ceiling'])} sqm")
        print(f"  Floor       : {fmt(tot['floor'])} sqm")
        print(f"  Total panel : {fmt(tot['total_panel'])} sqm")
        print(f"  Volume      : {fmt(tot['volume'])} m³")
        print(f"  Grand total : QAR {float(spec['total']):,.2f}")
        print(f"  In words    : {spec.get('total_words') or amount_in_words(spec['total'])}")


if __name__ == "__main__":
    main()
