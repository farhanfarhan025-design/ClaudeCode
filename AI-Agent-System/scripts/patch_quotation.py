#!/usr/bin/env python3
"""Apply job-specific overrides to a generated cold room quotation.

    python3 patch_quotation.py --docx quote.docx --overrides overrides.json

The `cold-room-quote` generator fills the fields that vary on every job —
client, dimensions, price, BOQ. Three things it treats as fixed boilerplate do
vary on some jobs, and stating them wrongly misdescribes the offer:

    flooring / refrigeration / capacity   when equipment is re-used rather
                                          than supplied, or the floor is not
                                          an insulated panel
    scope bullets                         when the job is a removal and
                                          re-installation, not a new build
    payment terms                         when the stages differ from the
                                          standard 75 / 20 / 5

This patches the produced document rather than the master, so the frozen
template — cover page, section photographs, brand styling — is never touched,
and the synced skill is left exactly as it is.

Overrides file:

    {
      "tables":  { "<table lookup label>": { "<row label>": "<value>" } },
      "bullets": { "8.  SCOPE OF WORK": ["...", "..."] }
    }

Tables are found by a label appearing in their first column, the same way the
generator finds them. Bullet sections are found by their heading paragraph, and
the run of paragraphs up to the next numbered heading is replaced.
"""

from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from pathlib import Path

from docx import Document

# A paragraph that ends a bullet run: the next numbered section, or the
# closing paragraph after the last section.
STOP = re.compile(r"^\d+\.\s+[A-Z]|^We trust|^Thank you for the opportunity")


def set_para_text(p, text: str) -> None:
    """Replace a paragraph's text, keeping the formatting of its first run.

    Writing to paragraph.text would drop the run properties and with them the
    bullet's font and colour."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def find_table(doc, label: str):
    for t in doc.tables:
        for row in t.rows:
            if row.cells and row.cells[0].text.strip() == label:
                return t
    raise LookupError(f"no table containing a row labelled {label!r}")


def set_row(table, label: str, value: str) -> None:
    for row in table.rows:
        cells = row.cells
        if cells and cells[0].text.strip() == label:
            # Merged layouts repeat the same cell object across the row; write
            # to each distinct cell after the label so the value is not left
            # showing the template's text in a duplicate column.
            seen = set()
            for cell in cells[1:]:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for i, p in enumerate(cell.paragraphs):
                    set_para_text(p, value if i == 0 else "")
            return
    raise LookupError(f"no row labelled {label!r}")


def replace_bullets(doc, heading: str, items: list[str]) -> int:
    """Replace the bullet run under `heading` with `items`.

    Surplus template bullets are deleted and extra ones cloned from the last,
    so the list length follows the job rather than the template."""
    paragraphs = doc.paragraphs
    start = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == heading.strip():
            start = i
            break
    if start is None:
        raise LookupError(f"heading {heading!r} not found")

    bullets = []
    for p in paragraphs[start + 1:]:
        text = p.text.strip()
        if text and STOP.match(text):
            break
        if text:
            bullets.append(p)

    if not bullets:
        raise LookupError(f"no bullets under {heading!r}")

    for i, text in enumerate(items):
        if i < len(bullets):
            set_para_text(bullets[i], text)
        else:
            new = copy.deepcopy(bullets[-1]._element)
            bullets[-1]._element.addnext(new)
            from docx.text.paragraph import Paragraph
            para = Paragraph(new, bullets[-1]._parent)
            set_para_text(para, text)
            bullets.append(para)

    for extra in bullets[len(items):]:
        extra._element.getparent().remove(extra._element)

    return len(items)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--docx", required=True, type=Path)
    ap.add_argument("--overrides", required=True, type=Path)
    args = ap.parse_args()

    spec = json.loads(args.overrides.read_text())
    doc = Document(str(args.docx))

    for label, rows in spec.get("tables", {}).items():
        try:
            table = find_table(doc, next(iter(rows)))
        except (LookupError, StopIteration):
            try:
                table = find_table(doc, label)
            except LookupError:
                print(f"  ! table {label!r} not found — skipped", file=sys.stderr)
                continue
        for row_label, value in rows.items():
            try:
                set_row(table, row_label, value)
                print(f"  table {label}: {row_label}")
            except LookupError:
                print(f"  ! row {row_label!r} not in {label!r} — skipped", file=sys.stderr)

    for heading, items in spec.get("bullets", {}).items():
        try:
            n = replace_bullets(doc, heading, items)
            print(f"  bullets {heading.strip()}: {n} items")
        except LookupError as exc:
            print(f"  ! {exc} — skipped", file=sys.stderr)

    doc.save(str(args.docx))
    print(f"\npatched → {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
