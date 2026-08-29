"""Unbranded Word versions of the Oscar Prime completion and handover certificates.

No letterhead, no company name anywhere — the top of page 1 is left clear so a letterhead can
be applied, and the entity above the signature is a blank rule to be completed by hand or by
whoever issues the document.
"""
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/tmp/claude-0/-home-user-ClaudeCode/f198eb8f-5169-5011-a0a8-bd4d87e7e71f/scratchpad/docx/"
NAVY = RGBColor(0x1F, 0x38, 0x64)


def new_doc(letterhead_cm=4.5):
    d = docx.Document()
    s = d.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.8)
    st = d.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.08
    # Clear space at the top of page 1 for a letterhead to be applied.
    gap = d.add_paragraph()
    gap.paragraph_format.space_after = Pt(0)
    gap.add_run().add_break()
    gap.paragraph_format.space_before = Cm(letterhead_cm)
    return d


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def title(d, text, sub=None):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    if sub:
        q = d.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = q.add_run(sub)
        rr.font.size = Pt(9.5)
        rr.italic = True
        q.paragraph_format.space_after = Pt(10)
    return p


def kv_table(d, rows, key_w=Cm(5.2)):
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        row = t.add_row()
        row.cells[0].width = key_w
        shade(row.cells[0], "EEF2F8")
        for cell, txt, bold in ((row.cells[0], k, True), (row.cells[1], v, True)):
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(1)
            par.paragraph_format.space_before = Pt(1)
            run = par.add_run(txt)
            run.bold = bold
            run.font.size = Pt(9.5)
            if cell is row.cells[0]:
                run.font.color.rgb = NAVY
    return t


def para(d, text, bold_prefix=None, size=10, after=6, justify=True):
    p = d.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def sig_block(d, right_label):
    d.add_paragraph().paragraph_format.space_after = Pt(10)
    t = d.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, label in zip(t.rows[0].cells, ("", right_label)):
        cell.width = Cm(8.4)
    left, right = t.rows[0].cells
    for cell, heading in ((left, None), (right, right_label)):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.add_run("_" * 40).font.size = Pt(10)
        if heading is None:
            q = cell.add_paragraph()
            q.paragraph_format.space_after = Pt(1)
            r = q.add_run("For and on behalf of")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            n = cell.add_paragraph()
            n.paragraph_format.space_after = Pt(1)
            n.add_run("_" * 40).font.size = Pt(10)
        else:
            q = cell.add_paragraph()
            q.paragraph_format.space_after = Pt(1)
            r = q.add_run(heading)
            r.bold = True
            r.font.size = Pt(9.5)
        f = cell.add_paragraph()
        f.paragraph_format.space_after = Pt(0)
        rr = f.add_run("Name / Signature / Date  ·  Company stamp")
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return t

# ============================================================ certificate of completion
d = new_doc()
title(d, "CERTIFICATE OF COMPLETION",
      "Supply, Installation, Testing and Commissioning of Cold Room")

kv_table(d, [
    ("Reference", "LTR/226/2026"),
    ("Date", "29 August 2026"),
])
d.add_paragraph().paragraph_format.space_after = Pt(4)

p = d.add_paragraph(); p.paragraph_format.space_after = Pt(1)
p.add_run("To,").font.size = Pt(10)
p = d.add_paragraph(); p.paragraph_format.space_after = Pt(0)
r = p.add_run("M/s. Oscar Prime Trading Contracting and Services W.L.L."); r.bold = True
p = d.add_paragraph(); p.paragraph_format.space_after = Pt(0)
p.add_run("P.O. Box 20545, Doha, State of Qatar").font.size = Pt(10)
p = d.add_paragraph(); p.paragraph_format.space_after = Pt(9)
p.add_run("Kind Attention: Mr. Shameem").font.size = Pt(10)

para(d, "", bold_prefix="SUBJECT: CERTIFICATE OF COMPLETION — SUPPLY, INSTALLATION, TESTING AND "
                        "COMMISSIONING OF COLD ROOM", after=9, justify=False)

para(d, "Dear Sir,", after=6, justify=False)
para(d, "We hereby certify that the works described below, carried out under your Local Purchase "
        "Order OTTS/LPO/19082026-02/2026, have been completed, tested and commissioned in "
        "accordance with the agreed scope of work and specification.", after=8)

kv_table(d, [
    ("Client", "Oscar Prime Trading Contracting and Services W.L.L."),
    ("Chiller Room 01", "2.59 m (L) × 1.80 m (W) × 3.20 m (H)"),
    ("Chiller Room 02", "3.94 m (L) × 2.31 m (W) × 3.20 m (H)"),
    ("Your Local Purchase Order", "OTTS/LPO/19082026-02/2026 dated 17-08-2026"),
    ("Date of Completion", "27 August 2026"),
    ("Date of Commissioning", "27 August 2026"),
    ("Handover Certificate", "HO/SQ074/2026 — signed at site by both parties"),
])
d.add_paragraph().paragraph_format.space_after = Pt(4)

para(d, "supply, fabrication and installation of two (2) cold rooms in 100 mm PUF sandwich panel "
        "to walls and ceiling, with aluminium coving, corner angles, capping, silicone sealing and "
        "accessories; two (2) hinged cold room doors, clear opening 900 × 1900 mm, 100 mm thick, "
        "with heavy-duty hinges, safety latch handle and internal safety release; the refrigeration "
        "system — BITZER (Germany) semi-hermetic condensing units with matching Guntner (Germany) "
        "evaporating units, refrigeration copper piping, Armaflex insulation and insulated drain "
        "lines; the digital control panel with all safeties, internal LED lighting, door frame "
        "heaters and the electrical works from the panel onward; and pressure testing, evacuation, "
        "gas charging, testing and commissioning, transportation, mobilisation and lifting at site, "
        "with handover and operation training to your staff.",
     bold_prefix="The completed works comprise: ", after=8)

box = d.add_table(rows=1, cols=1); box.style = "Table Grid"
shade(box.rows[0].cells[0], "EEF2F8")
c = box.rows[0].cells[0]
q = c.paragraphs[0]; q.paragraph_format.space_after = Pt(3)
r = q.add_run("Warranty. "); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(9.5)
q.add_run("Twelve (12) months on the panels and doors from the date of handover, and twelve (12) "
          "months on the compressors, coils, electrical components, controls and workmanship from "
          "the date of commissioning. Excludes damage from misuse, power surge, improper or "
          "interrupted supply, water ingress, unauthorised modification and absence of routine "
          "maintenance.").font.size = Pt(9.5)
q2 = c.add_paragraph(); q2.paragraph_format.space_after = Pt(1)
r = q2.add_run("Design basis. "); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(9.5)
q2.add_run("The rooms are designed and commissioned for chiller duty at 0 °C to +5 °C. Operation "
           "outside that range is outside the design basis and is not covered by the warranty "
           "above.").font.size = Pt(9.5)
d.add_paragraph().paragraph_format.space_after = Pt(4)

para(d, "The verification of every element is recorded on the Handover Certificate HO/SQ074/2026, "
        "signed jointly at site. We thank you for the opportunity to carry out this project and "
        "remain at your service for maintenance and any future requirement.", after=8)
para(d, "Yours faithfully,", after=2, justify=False)

sig_block(d, "Received and accepted — Oscar Prime Trading Cont. & Services W.L.L.")
d.save(OUT + "Certificate_of_Completion_OscarPrime.docx")
print("completion saved")


# ============================================================ handover certificate
import json

hj = json.load(open(SPJSON := "/tmp/claude-0/-home-user-ClaudeCode/"
                              "f198eb8f-5169-5011-a0a8-bd4d87e7e71f/scratchpad/issue/"
                              "oscar_handover.json"))

d = new_doc()
title(d, "HANDOVER CERTIFICATE",
      "Supply, Installation, Testing and Commissioning of Cold Room")

kv_table(d, [
    ("Certificate No.", "HO/SQ074/2026"),
    ("Date", "27-08-2026"),
    ("Reference", "LPO OTTS/LPO/19082026-02/2026"),
])
d.add_paragraph().paragraph_format.space_after = Pt(4)

p = d.add_paragraph(); p.paragraph_format.space_after = Pt(1)
r = p.add_run("HANDOVER TO"); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY

kv_table(d, [
    ("Client", "Oscar Prime Trading Contracting and Services W.L.L."),
    ("Address", "P.O. Box 20545, Doha, State of Qatar"),
    ("Kind Attention", "Mr. Shameem"),
    ("Chiller Room 01", "2.59 m (L) × 1.80 m (W) × 3.20 m (H)"),
    ("Chiller Room 02", "3.94 m (L) × 2.31 m (W) × 3.20 m (H)"),
    ("Date of Completion", "27-08-2026"),
    ("Date of Commissioning", "27-08-2026"),
], key_w=Cm(4.4))
d.add_paragraph().paragraph_format.space_after = Pt(6)

para(d, "The following items have been jointly inspected at site and verified as complete.",
     after=5, justify=False)

WIDTHS = (Cm(1.1), Cm(10.9), Cm(1.9), Cm(1.3), Cm(1.8))
HEADS = ("S/N", "DESCRIPTION OF WORK VERIFIED", "UNIT", "QTY", "VERIFIED")

t = d.add_table(rows=1, cols=5)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
hdr = t.rows[0]
for i, (cell, head) in enumerate(zip(hdr.cells, HEADS)):
    cell.width = WIDTHS[i]
    shade(cell, "1F3864")
    par = cell.paragraphs[0]
    par.paragraph_format.space_before = Pt(2)
    par.paragraph_format.space_after = Pt(2)
    if i != 1:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(head)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Repeat the header row at the top of each page.
trPr = hdr._tr.get_or_add_trPr()
th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)

for n, line in enumerate(hj["lines"], 1):
    desc = line["description"]
    row = t.add_row()
    vals = (str(n), desc, line["unit"], str(line["qty"]), "☐")
    for i, (cell, txt) in enumerate(zip(row.cells, vals)):
        cell.width = WIDTHS[i]
        par = cell.paragraphs[0]
        par.paragraph_format.space_before = Pt(1.5)
        par.paragraph_format.space_after = Pt(1.5)
        if i != 1:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 1 and " — " in txt and txt.split(" — ")[0].isupper():
            head, rest = txt.split(" — ", 1)
            rh = par.add_run(head + " — "); rh.bold = True; rh.font.size = Pt(8.5)
            rb = par.add_run(rest); rb.font.size = Pt(8.5)
        else:
            run = par.add_run(txt)
            run.font.size = Pt(9) if i == 4 else Pt(8.5)

d.add_paragraph().paragraph_format.space_after = Pt(6)

p = d.add_paragraph(); p.paragraph_format.space_after = Pt(3)
r = p.add_run("NOTES"); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = NAVY

for note in hj["notes"]:
    np_ = d.add_paragraph(style="List Number")
    np_.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    np_.paragraph_format.space_after = Pt(4)
    if " — " in note and note.split(" — ")[0].isupper():
        head, rest = note.split(" — ", 1)
        rh = np_.add_run(head + " — "); rh.bold = True; rh.font.size = Pt(9)
        rb = np_.add_run(rest); rb.font.size = Pt(9)
    else:
        np_.add_run(note).font.size = Pt(9)

sig_block(d, "Received and accepted — Oscar Prime Trading Cont. & Services W.L.L.")
d.save(OUT + "Handover_Certificate_OscarPrime.docx")
print("handover saved")
