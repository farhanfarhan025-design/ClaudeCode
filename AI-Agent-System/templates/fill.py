import copy, docx
from docx.shared import Pt
from docx.oxml.ns import qn

SRC = "/root/.claude/uploads/f198eb8f-5169-5011-a0a8-bd4d87e7e71f/"
OUT = "/tmp/claude-0/-home-user-ClaudeCode/f198eb8f-5169-5011-a0a8-bd4d87e7e71f/scratchpad/hse/"

def setcell(cell, text, bold=False, size=None):
    """Write text into a cell, keeping the template's own font where one exists."""
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    model = para.runs[0] if para.runs else None
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    lines = str(text).split("\n")
    for i, ln in enumerate(lines):
        if i:
            para.add_run().add_break()
        run = para.add_run(ln)
        if model is not None:
            run.font.name = model.font.name
            run.font.size = model.font.size
            rpr = run._element.get_or_add_rPr()
            rf = model._element.find(qn('w:rPr'))
            if rf is not None and rf.find(qn('w:rFonts')) is not None:
                rpr.append(copy.deepcopy(rf.find(qn('w:rFonts'))))
        if size:
            run.font.size = Pt(size)
        run.bold = bold


def append_text(cell, text, size=None):
    """Add a paragraph to a cell that already has a label in it."""
    para = cell.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size or 9)
    return para

# ===================================================================== JHA
d = docx.Document(SRC + "dc544bf1-JOB_HAZARD_ANALYSIS_GWCHSEFMJHA.00.docx")
t = d.tables[0]

def C(r, c):
    return t.rows[r].cells[c]

setcell(C(2, 0), "CLIENT/ CONTRACTOR NAME:  The New Doha Kitchen Equipment Services W.L.L. (TNDK)")
setcell(C(2, 6), "SITE:  Al Zehrabi Medical L.L.C. — cold room area and adjacent outdoor plant "
                 "area, B11 9")
setcell(C(3, 0), "DESCRIPTION OF THE ACTIVITY:  Cold room maintenance — water servicing of two (2) "
                 "refrigeration units (condensing unit and evaporator each); nitrogen pressure test and "
                 "brazing repair of a refrigerant leak on one (1) cold room circuit; evacuation, refrigerant "
                 "charging, testing and commissioning; removal and renewal of silicone sealant to panel "
                 "joints; cold room door adjustment; installation of 25 m of pipe insulation; and "
                 "fabrication and installation of a 3.00 m x 1.50 m sunshade over the outdoor condensing "
                 "units. Reference: Quotation QUT/DCTS/219/2026 and Work Order WO/DCTS/001/2026.")
setcell(C(4, 0), "Prepared by:  Farhan — Sales Engineer, TNDK        Date: 19-08-2026\n(Name & Signature)")
setcell(C(4, 6), "Reviewed and Approved by:\n(Name & Signature)")

for row, col in ((6, 0), (7, 0), (8, 0), (6, 4), (7, 4), (8, 4)):
    setcell(C(row, col), "X", bold=True)
setcell(C(6, 9), "Others: (Please specify)  Face shield and leather gauntlets for brazing; "
                 "cut-resistant gloves for sheet steel; insulated gloves and thermal jacket for work "
                 "inside the cold room; dust mask when drilling.")

HAZARDS = [
 ("Mobilisation, site set-up and moving materials, cylinders and tools to the work area",
  "Manual handling of gas cylinders, steel section and insulation. Slips, trips and falls from hoses, cables and wet floor. Vehicle movement in the yard",
  "Back and muscle injury. Crush injury to hands and feet. Falls causing fracture or sprain. Being struck by a vehicle",
  "Team lift or trolley for anything over 25 kg. Cylinders moved upright on a trolley, secured, valve caps on. Work area barricaded and signed. Hoses and cables routed clear of walkways and covered where they cross. Spillages mopped immediately. High-visibility vest worn at all times in the yard",
  "Site Supervisor"),
 ("Electrical isolation before servicing and opening any panel",
  "Live electrical parts, stored energy in capacitors, unexpected re-energisation by others",
  "Electric shock, burns, arc flash, fatality",
  "Isolate at source, lock off and tag out, and keep the key with the technician doing the work. Prove dead with a tested meter before touching any conductor. Only a competent person opens a panel. Insulated tools. No isolation to be removed by anyone other than the person who applied it",
  "Refrigeration Technician"),
 ("Water washing of condenser and evaporator coils under pressure",
  "Water on and near live electrical equipment. High-pressure water jet. Wet, slippery floor. Noise from the washer",
  "Electric shock. Injury to eyes and skin from the jet. Slips and falls. Hearing damage",
  "Power isolated and locked off before any water is introduced. Motors, terminal boxes and controls covered before washing. Jet never directed at electrical parts or at another person. Safety glasses worn. Ear protection when the washer is running. Floor squeegeed and dried before power is restored, and the area left dry before leaving",
  "Refrigeration Technician"),
 ("Refrigerant recovery and handling",
  "Pressurised refrigerant. Liquid refrigerant contact with skin or eyes. Vapour displacing oxygen in an enclosed cold room",
  "Frostbite and cold burns. Eye injury. Asphyxiation in a confined space",
  "Refrigerant recovered to a recovery machine and cylinder — never vented to atmosphere. Gloves and safety glasses worn when breaking into any circuit. Cold room door held open and the space ventilated while the circuit is open. Nobody works alone inside the room. Recovery cylinder weighed and not overfilled",
  "Refrigeration Technician"),
 ("Nitrogen pressure testing of the refrigerant circuit",
  "High-pressure nitrogen. Over-pressurisation of the system. Cylinder falling or being struck. Uncontrolled release of a hose or fitting",
  "Explosive rupture of pipework or components. Flying debris and eye injury. Impact injury from a falling cylinder. Asphyxiation from nitrogen build-up in an enclosed room",
  "Nitrogen introduced ONLY through a regulator with a relief valve — never direct from the cylinder. Test pressure limited to the manufacturer's stated maximum for the system and recorded on the work order. Cylinder secured upright and chained. All personnel clear during pressurisation. Hoses and fittings checked before use. Room ventilated throughout. Safety glasses worn",
  "Refrigeration Technician"),
 ("Brazing repair of the leak (hot work)",
  "Naked flame and hot surfaces. Oxy-fuel and MAPP gas cylinders. Combustible material — panel insulation, packaging, insulation foam. Fumes from brazing and from burning insulation",
  "Fire and spread of fire through panel insulation. Burns. Damage to the cold room fabric. Smoke inhalation",
  "HOT WORK PERMIT to be obtained before any flame is lit. Combustible material removed for 10 m or protected with a fire blanket. Fire extinguisher at the work position and a fire watch kept during the work and for 30 minutes after it. Nitrogen purge while brazing. NEVER braze a circuit that still holds refrigerant — recover first. Face shield and leather gauntlets worn. Cylinders upright, secured and outside the room where practical",
  "Refrigeration Technician / Fire Watch"),
 ("Evacuation, gas charging, testing and commissioning",
  "Pressurised system. Rotating fans and belts on start-up. Refrigerant release at a hose connection",
  "Frostbite, eye injury, entanglement, cuts to hands",
  "Guards refitted before any unit is started. Nobody near fan blades on start-up. Gauge hoses and connections checked before charging. Charge by weight to the nameplate figure and record it. Gloves and safety glasses worn throughout",
  "Refrigeration Technician"),
 ("Working inside the cold room — silicone removal and renewal, door adjustment",
  "Entrapment inside the room if the door closes. Cold stress. Sharp knives and scrapers for raking out old sealant. Solvents, sealant and cleaning chemicals",
  "Person shut in a cold room. Hypothermia. Cuts to hands. Skin and eye irritation, fumes in an enclosed space",
  "CONFIRM THE INTERNAL SAFETY RELEASE WORKS BEFORE ANYONE ENTERS. Door wedged open while work is in progress, and never allowed to latch with a person inside. Nobody works alone in the room — a second person stays outside and keeps contact. Cutting stroke always away from the body, cut-resistant gloves worn. Chemicals used with the door open and the space ventilated, gloves and safety glasses worn, and the safety data sheet kept on site",
  "Site Supervisor"),
 ("Work at height — ceiling evaporators, pipe insulation runs and sunshade installation",
  "Falls from a ladder or platform. Tools and materials dropped on people below. Overreaching",
  "Fracture and serious injury from a fall. Head injury from dropped tools",
  "Ladder inspected before use, set at the correct angle, footed or tied, and used only for light short-duration work — a platform or tower for anything longer. Three points of contact maintained; no overreaching, move the ladder instead. Area below barricaded while overhead work is in progress. Tools carried in a belt or raised by line, never thrown. Hard hats worn by everyone in the area",
  "Site Supervisor"),
 ("Fabrication and installation of the sunshade — cutting, drilling, grinding and painting steel",
  "Sparks and hot metal. Sharp edges and swarf. Flying particles from drilling and grinding. Noise. Paint fumes. Weight of the fabricated frame during lifting into position",
  "Eye injury. Cuts and lacerations. Burns. Hearing damage. Inhalation of fumes. Crush injury while positioning the frame",
  "Grinding and drilling only with the guard fitted and with safety glasses or a face shield. Cut edges deburred before handling; cut-resistant gloves worn. Sparks kept away from the cold room panels and any combustible material, and a hot work permit obtained if grinding near them. Painting in the open air, away from air intakes. Frame lifted into position by a minimum of two persons, or mechanically. Clearance to the condenser air intake and discharge checked and confirmed before the frame is finally fixed",
  "Site Supervisor"),
 ("Housekeeping and handover on completion",
  "Waste sealant, offcuts, packaging and old insulation left on site. Tools and cylinders left unattended",
  "Slips and trips. Cuts from sharp offcuts. Unauthorised person interfering with equipment",
  "Work area cleared progressively, not only at the end. All waste removed from site by TNDK. Cylinders removed or secured. Guards, panels and covers refitted and all isolations removed only when the work is complete. Area handed back clean and the client's representative walked through it before the team leaves",
  "Site Supervisor"),
]

blank = list(range(11, 17))
while len(blank) < len(HAZARDS):
    new = copy.deepcopy(t.rows[16]._tr)
    t._tbl.append(new)
    blank.append(len(t.rows) - 1)

for i, (act, haz, eff, ctl, resp) in enumerate(HAZARDS):
    r = blank[i]
    setcell(C(r, 0), act, size=8)
    setcell(C(r, 2), haz, size=8)
    setcell(C(r, 4), eff, size=8)
    setcell(C(r, 7), ctl, size=8)
    setcell(C(r, 11), resp, size=8)

d.save(OUT + "JHA_AlZehrabi_ColdRoomMaintenance_GWC-HSE-FM-JHA.00.docx")
print("JHA saved — hazard rows:", len(HAZARDS))
