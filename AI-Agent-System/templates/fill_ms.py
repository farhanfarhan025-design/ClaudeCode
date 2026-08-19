import sys
sys.path.insert(0, "/tmp/claude-0/-home-user-ClaudeCode/f198eb8f-5169-5011-a0a8-bd4d87e7e71f/scratchpad/hse")
import copy, docx
from docx.shared import Pt
from docx.oxml.ns import qn

SRC = "/root/.claude/uploads/f198eb8f-5169-5011-a0a8-bd4d87e7e71f/"
OUT = "/tmp/claude-0/-home-user-ClaudeCode/f198eb8f-5169-5011-a0a8-bd4d87e7e71f/scratchpad/hse/"

def setcell(cell, text, bold=False, size=None):
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    model = para.runs[0] if para.runs else None
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for i, ln in enumerate(str(text).split("\n")):
        if i:
            para.add_run().add_break()
        run = para.add_run(ln)
        if model is not None:
            run.font.name = model.font.name
            run.font.size = model.font.size
            rf = model._element.find(qn('w:rPr'))
            if rf is not None and rf.find(qn('w:rFonts')) is not None:
                run._element.get_or_add_rPr().append(copy.deepcopy(rf.find(qn('w:rFonts'))))
        if size:
            run.font.size = Pt(size)
        run.bold = bold

d = docx.Document(SRC + "6edb9208-Method_Statement_GWCHSEFM11MS.00_1_Sample.docx")
t0, t1, t2 = d.tables[0], d.tables[1], d.tables[2]

def R(t, r):
    seen, out = [], []
    for c in t.rows[r].cells:
        if any(id(c._tc) == id(x._tc) for x in seen):
            continue
        seen.append(c); out.append(c)
    return out

setcell(R(t0, 2)[1], "The New Doha Kitchen Equipment Services W.L.L. (TNDK)")
setcell(R(t0, 3)[1], "Cold room maintenance — servicing of two refrigeration units, nitrogen leak test "
                     "and brazing repair of one circuit, re-gassing and commissioning, silicone sealing, "
                     "cold room door adjustment, 25 m of pipe insulation, and installation of a "
                     "3.00 m x 1.50 m sunshade over the outdoor condensing units.")
row4 = R(t0, 4)
setcell(row4[1], "[ GWC facility / area — to be inserted ]")
setcell(row4[3], "Al Zehrabi Medical L.L.C. — cold room area and adjacent outdoor plant area "
                 "[ building / unit no. to be inserted ]")
setcell(R(t0, 7)[0], "Cold room maintenance: servicing, leak rectification, re-gassing, sealing, door "
                     "adjustment, pipe insulation and outdoor sunshade installation.")

TOOLS = [
 ("Refrigerant recovery machine and recovery cylinder", "Vacuum pump and micron gauge", "Gauge manifold and charging hoses"),
 ("Dry nitrogen cylinder with regulator and relief valve", "Oxy-MAPP brazing set with flashback arrestors", "Electronic leak detector and soap solution"),
 ("Pressure washer, coil cleaner and fin comb", "Digital clamp meter and thermometer", "Silicone gun, rake and cleaning solvent"),
 ("Angle grinder, drill and hand tools (110 V / battery)", "Step ladder and mobile platform", "Fire extinguisher and fire blanket"),
 ("Refrigerant to suit the units, filter drier, insulation, sealant, steel section and fixings", "Barricade tape, cones and signage", "PPE as listed in the JHA"),
]
for i, trio in enumerate(TOOLS):
    cells = R(t0, 9 + i)
    for j, txt in enumerate(trio):
        if j < len(cells):
            setcell(cells[j], txt, size=9)

setcell(R(t0, 15)[0],
 "A work permit and, for the brazing repair, a HOT WORK PERMIT shall be obtained from GWC / the "
 "facility management before any work begins.\n"
 "This method statement and the attached Job Hazard Analysis will be discussed with every member "
 "of the site team, and a toolbox talk recorded before work starts each day.\n"
 "The work area — inside the cold room and around the outdoor units — will be barricaded with tape "
 "and cones, and warning signage displayed. Access will be restricted to the working team.\n"
 "The client will be asked to empty the cold rooms of stock, or relocate it, and to confirm in "
 "writing that the rooms may be taken out of service.\n"
 "The internal safety release on the cold room door will be tested before any person enters the "
 "room, and the door will be wedged open while work is in progress.\n"
 "Electrical supply to the units will be isolated, locked off and tagged, and proved dead before "
 "any panel is opened or any washing begins.\n"
 "Combustible material will be removed from the brazing area or protected with a fire blanket, and "
 "a fire extinguisher with a nominated fire watch positioned at the work point.\n"
 "Gas cylinders will be transported and stored upright, secured against falling, with valve caps "
 "fitted when not in use.\n"
 "Emergency contact numbers and the route to the nearest muster point will be confirmed with the "
 "facility before work starts.", size=9)

setcell(R(t0, 17)[0],
 "1. TOOLBOX TALK AND PERMITS. The team attends a toolbox talk covering the scope, the hazards in "
 "the attached JHA and the emergency arrangements. Permits are displayed at the work area.\n"
 "2. PRE-WORK RECORD. Existing condition photographed, nameplate details and running readings "
 "recorded on Work Order WO/DCTS/001/2026 before any work begins.\n"
 "3. ISOLATION. Power isolated at source, locked off and tagged, and proved dead. The key is held "
 "by the technician carrying out the work.\n"
 "4. WATER SERVICE OF TWO UNITS. Electrical parts covered. Condenser and evaporator coils washed "
 "with approved cleaner and pressurised water, fins combed, fans and guards cleaned, drain trays "
 "and lines flushed and proved clear. Terminals checked and tightened. Floor dried before power is "
 "restored, then units run and readings recorded.\n"
 "5. RECOVERY AND LEAK TEST. Remaining refrigerant recovered to a cylinder — never vented. The "
 "circuit is pressurised with dry nitrogen through a regulator, held under observation, and the "
 "leak traced with leak detector and soap solution. Pressure and holding time are recorded.\n"
 "6. STOP POINT. If the leak is found inside an evaporator or condenser coil, or within the "
 "compressor body, work STOPS. No repair is attempted, the finding is reported with the test "
 "result, and the replacement is quoted separately before any further work.\n"
 "7. BRAZING REPAIR. Carried out under hot work permit with a fire watch and extinguisher in place, "
 "combustibles cleared or protected, and a nitrogen purge maintained. The circuit is then "
 "re-pressurised with nitrogen and held a second time to prove it is holding before any refrigerant "
 "is introduced.\n"
 "8. EVACUATION AND CHARGING. A new filter drier is fitted, the circuit evacuated to a deep vacuum "
 "and a vacuum hold test carried out and recorded. Refrigerant is charged by weight to the unit "
 "nameplate, and the type and weight recorded.\n"
 "9. TESTING AND COMMISSIONING. Guards refitted before start-up. Suction and discharge pressures, "
 "superheat, compressor current and room pull-down are run and recorded.\n"
 "10. SILICONE WORK. Failed sealant raked out at panel joints and wall-to-ceiling and wall-to-floor "
 "junctions, surfaces cleaned and dried, and re-sealed in food-grade anti-fungal silicone. The room "
 "is ventilated and the door held open throughout, with a second person outside.\n"
 "11. DOOR ADJUSTMENT. Hinges and closer adjusted, the leaf realigned and squared, the gasket "
 "seated, and the seal proved by a light test from inside the room.\n"
 "12. PIPE INSULATION. 25 m of closed-cell insulation installed on the refrigerant lines, all joints "
 "glued and sealed, and exposed sections given a protective finish.\n"
 "13. SUNSHADE. The steel frame is cut, drilled, deburred, treated and painted, then lifted into "
 "position by not fewer than two persons and fixed over the outdoor units. Clearance to the "
 "condenser air intake and discharge is checked and confirmed before final fixing.\n"
 "14. HOUSEKEEPING AND HANDOVER. Guards and covers refitted, isolations removed, all waste removed "
 "from site, final readings recorded, completed work photographed, and the area walked through with "
 "the client's representative, who signs the work order.", size=9)

setcell(R(t0, 19)[0], "No")

def fill_person(row, name, desig, mob, mail):
    cells = R(t1, row)
    for c, v in zip(cells, (name, desig, mob, mail)):
        setcell(c, v, size=9)

# r02 carries the template's sample person — Ashley Fernandes of another company. It must be
# overwritten, not left behind: submitting a form with a stranger's name, number and email on it
# is both a data-protection problem and an obvious sign the form was not read.
fill_person(2, "Joseph Jobi", "Engineer — Key Personnel / Site in charge", "7742 2762",
            "farhan@dctsqatar.com")
fill_person(3, "Jijo Maliyekkal Thomas\nQID 28035608170 (exp. 10-11-2026)",
            "Refrigeration Technician", "[ mobile ]", "")
fill_person(4, "Md Mijan Molla Yeleus Mollah\nQID 29805016707 (exp. 30-12-2026)",
            "Refrigeration Technician", "[ mobile ]", "")
fill_person(5, "Tauhid Hossen Ashrafujjaman\nQID 30605005039 (exp. 02-10-2026)",
            "Helper", "[ mobile ]", "")

# r06 and r08 are the template's own "Work Force" and "Additional Documents" labels — the text
# goes in the blank row under each, not over the label.
setcell(R(t1, 7)[0], "1 x Engineer / Site in charge, 2 x Refrigeration Technicians, 1 x Helper — "
                     "4 persons in total. QID copies attached for the three site staff. "
                     "[ Attach the Engineer's QID copy and insert mobile numbers before submission. ]", size=9)
setcell(R(t1, 9)[0], "Quotation QUT/DCTS/219/2026 (approved)  ·  Work Order "
                     "WO/DCTS/001/2026  ·  Job Hazard Analysis GWC-HSE-FM-JHA.00 (attached)  ·  Hot Work "
                     "Permit (to be obtained)  ·  Company trade licence and CR  ·  Third-party liability "
                     "insurance certificate  ·  QID and gate-pass copies for all personnel  ·  Gas cylinder "
                     "test certificates  ·  Safety data sheets for refrigerant, silicone and solvents.", size=9)

setcell(R(t2, 1)[0], "Name:  Farhan")
setcell(R(t0, 6)[0], "Work Title:")
setcell(R(t2, 2)[0], "Designation:  Sales Engineer")
setcell(R(t2, 4)[0], "Date:  19-08-2026")

d.save(OUT + "MethodStatement_AlZehrabi_ColdRoomMaintenance_GWC-HSE-FM11-MS.00.docx")
print("Method statement saved")
